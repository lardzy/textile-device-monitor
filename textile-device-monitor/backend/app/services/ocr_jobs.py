from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime, timedelta, timezone
import io
import json
from pathlib import Path
import queue
import shutil
import threading
from typing import Any
from uuid import uuid4

from docx import Document
from docx.shared import Inches
from PIL import Image
import pypdfium2 as pdfium
import requests

from app.config import settings


class OcrJobError(Exception):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code
        self.message = message


@dataclass
class OcrJobRecord:
    job_id: str
    original_filename: str
    upload_path: str
    source_path: str
    output_dir: str
    markdown_path: str
    json_path: str
    page_range: str | None
    note: str | None
    status: str = "queued"
    created_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    started_at: datetime | None = None
    finished_at: datetime | None = None
    error_code: str | None = None
    error_message: str | None = None


class OcrJobManager:
    def __init__(self) -> None:
        self._queue: queue.Queue[str] = queue.Queue()
        self._jobs: dict[str, OcrJobRecord] = {}
        self._lock = threading.Lock()
        self._shutdown_event = threading.Event()
        self._workers: list[threading.Thread] = []
        self._started = False

    def start(self) -> None:
        if not settings.OCR_ENABLED:
            return
        with self._lock:
            if self._started:
                return
            Path(settings.OCR_UPLOAD_DIR).mkdir(parents=True, exist_ok=True)
            Path(settings.OCR_OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
            worker_count = max(1, int(settings.OCR_MAX_CONCURRENT_JOBS))
            self._shutdown_event.clear()
            self._workers = []
            for index in range(worker_count):
                worker = threading.Thread(
                    target=self._worker_loop,
                    name=f"ocr-worker-{index + 1}",
                    daemon=True,
                )
                worker.start()
                self._workers.append(worker)
            self._started = True

    def stop(self) -> None:
        with self._lock:
            if not self._started:
                return
            self._shutdown_event.set()
            workers = list(self._workers)
            self._workers = []
            self._started = False

        for worker in workers:
            worker.join(timeout=1.5)

    def create_job(
        self,
        upload_path: str,
        original_filename: str,
        page_range: str | None = None,
        note: str | None = None,
    ) -> dict[str, Any]:
        self.start()

        job_id = uuid4().hex
        output_dir = Path(settings.OCR_OUTPUT_DIR) / job_id
        output_dir.mkdir(parents=True, exist_ok=True)

        record = OcrJobRecord(
            job_id=job_id,
            original_filename=original_filename,
            upload_path=upload_path,
            source_path=str(
                output_dir / f"source{Path(original_filename).suffix.lower()}"
            ),
            output_dir=str(output_dir),
            markdown_path=str(output_dir / "result.md"),
            json_path=str(output_dir / "result.json"),
            page_range=(page_range or "").strip() or None,
            note=(note or "").strip() or None,
        )
        try:
            shutil.copy2(upload_path, record.source_path)
        except Exception:
            record.source_path = ""
        with self._lock:
            self._jobs[job_id] = record
        self._queue.put(job_id)
        payload = self.get_job(job_id)
        if payload is None:
            raise RuntimeError("failed_to_create_job")
        return payload

    def get_job(self, job_id: str) -> dict[str, Any] | None:
        with self._lock:
            record = self._jobs.get(job_id)
            if record is None:
                return None
            queue_position = self._get_queue_position_unlocked(job_id)
            return self._serialize_job(record, queue_position)

    def get_result(self, job_id: str) -> dict[str, Any] | None:
        with self._lock:
            record = self._jobs.get(job_id)
            if record is None or record.status != "succeeded":
                return None
            markdown_path = Path(record.markdown_path)
            json_path = Path(record.json_path)

        if not markdown_path.exists() or not json_path.exists():
            return None

        markdown_text = markdown_path.read_text(encoding="utf-8")
        with json_path.open("r", encoding="utf-8") as handle:
            json_data = json.load(handle)
        return {
            "job_id": job_id,
            "markdown_text": markdown_text,
            "json_data": json_data,
            "artifacts": {
                "md": {"filename": markdown_path.name},
                "json": {"filename": json_path.name},
            },
        }

    def get_artifact_path(self, job_id: str, kind: str) -> Path | None:
        with self._lock:
            record = self._jobs.get(job_id)
            if record is None or record.status != "succeeded":
                return None
            path = (
                Path(record.markdown_path)
                if kind == "md"
                else Path(record.json_path)
                if kind == "json"
                else None
            )
        if path is None or not path.exists():
            return None
        return path

    def cleanup_expired(self, retention_hours: int) -> int:
        cutoff = datetime.now(timezone.utc) - timedelta(hours=retention_hours)
        with self._lock:
            expired_job_ids = [
                job_id
                for job_id, job in self._jobs.items()
                if job.status in {"succeeded", "failed"}
                and job.finished_at is not None
                and job.finished_at < cutoff
            ]
            for job_id in expired_job_ids:
                self._jobs.pop(job_id, None)

        deleted_count = 0
        for job_id in expired_job_ids:
            job_output_dir = Path(settings.OCR_OUTPUT_DIR) / job_id
            if job_output_dir.exists():
                shutil.rmtree(job_output_dir, ignore_errors=True)
            deleted_count += 1
        return deleted_count

    def _get_queue_position_unlocked(self, job_id: str) -> int | None:
        queued = sorted(
            (
                (record.created_at, queued_job_id)
                for queued_job_id, record in self._jobs.items()
                if record.status == "queued"
            ),
            key=lambda item: item[0],
        )
        for index, (_, queued_job_id) in enumerate(queued, start=1):
            if queued_job_id == job_id:
                return index
        return None

    def _serialize_job(
        self, record: OcrJobRecord, queue_position: int | None = None
    ) -> dict[str, Any]:
        return {
            "job_id": record.job_id,
            "status": record.status,
            "created_at": record.created_at.isoformat(),
            "started_at": (
                record.started_at.isoformat() if record.started_at is not None else None
            ),
            "finished_at": (
                record.finished_at.isoformat()
                if record.finished_at is not None
                else None
            ),
            "error_code": record.error_code,
            "error_message": record.error_message,
            "queue_position": queue_position,
            "original_filename": record.original_filename,
        }

    def export_jobs_docx(self, job_ids: list[str]) -> tuple[Path, str]:
        normalized_ids = [
            job_id.strip() for job_id in job_ids if job_id and job_id.strip()
        ]
        if not normalized_ids:
            raise ValueError("empty_job_ids")

        records: list[OcrJobRecord] = []
        with self._lock:
            for job_id in normalized_ids:
                record = self._jobs.get(job_id)
                if record is None:
                    raise ValueError("job_not_found")
                if record.status != "succeeded":
                    raise ValueError("job_not_completed")
                records.append(record)

        export_dir = Path(settings.OCR_OUTPUT_DIR) / "_exports"
        export_dir.mkdir(parents=True, exist_ok=True)
        filename = (
            f"ocr-batch-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S')}.docx"
        )
        export_path = export_dir / f"{uuid4().hex}.docx"

        document = Document()

        for index, record in enumerate(records, start=1):
            if index > 1:
                document.add_page_break()

            markdown_path = Path(record.markdown_path)
            json_path = Path(record.json_path)
            if not markdown_path.exists() or not json_path.exists():
                continue

            markdown_text = markdown_path.read_text(encoding="utf-8")
            for line in markdown_text.splitlines():
                document.add_paragraph(line)

            with json_path.open("r", encoding="utf-8") as handle:
                json_data = json.load(handle)

            illustration_chunks = self._extract_illustration_chunks(
                json_data=json_data,
                source_path=Path(record.source_path) if record.source_path else None,
            )
            if illustration_chunks:
                if markdown_text.strip():
                    document.add_paragraph("")
                for chunk in illustration_chunks:
                    image_stream = io.BytesIO(chunk["image_bytes"])
                    image_stream.seek(0)
                    document.add_picture(image_stream, width=Inches(5.8))

        document.save(export_path)
        return export_path, filename

    def _extract_illustration_chunks(
        self,
        json_data: Any,
        source_path: Path | None,
    ) -> list[dict[str, Any]]:
        if source_path is None or not source_path.exists():
            return []

        regions = _collect_image_regions(json_data)
        if not regions:
            return []

        try:
            pages = _render_source_pages(
                source_path, sorted({r["page_number"] for r in regions})
            )
        except Exception:
            return []

        chunks: list[dict[str, Any]] = []
        try:
            for region in regions:
                page_image = pages.get(region["page_number"])
                if page_image is None:
                    continue

                bbox = region["bbox"]
                x1, y1, x2, y2 = _sanitize_bbox(
                    bbox, page_image.width, page_image.height
                )
                if x2 - x1 < 12 or y2 - y1 < 12:
                    continue

                crop = page_image.crop((x1, y1, x2, y2))
                buf = io.BytesIO()
                crop.save(buf, format="PNG")
                chunks.append(
                    {
                        "page_number": region["page_number"],
                        "label": region["label"],
                        "image_bytes": buf.getvalue(),
                    }
                )
            return chunks
        finally:
            for image in pages.values():
                image.close()

    def _worker_loop(self) -> None:
        while not self._shutdown_event.is_set():
            try:
                job_id = self._queue.get(timeout=0.5)
            except queue.Empty:
                continue
            try:
                self._run_job(job_id)
            finally:
                self._queue.task_done()

    def _run_job(self, job_id: str) -> None:
        with self._lock:
            record = self._jobs.get(job_id)
            if record is None:
                return
            record.status = "running"
            record.started_at = datetime.now(timezone.utc)
            record.finished_at = None
            record.error_code = None
            record.error_message = None

        try:
            markdown_text, json_data = self._call_ocr_service(record)
            Path(record.markdown_path).write_text(markdown_text, encoding="utf-8")
            Path(record.json_path).write_text(
                json.dumps(json_data, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            with self._lock:
                current = self._jobs.get(job_id)
                if current is None:
                    return
                current.status = "succeeded"
                current.finished_at = datetime.now(timezone.utc)
                current.error_code = None
                current.error_message = None
        except OcrJobError as exc:
            with self._lock:
                current = self._jobs.get(job_id)
                if current is None:
                    return
                current.status = "failed"
                current.finished_at = datetime.now(timezone.utc)
                current.error_code = exc.code
                current.error_message = exc.message
        except Exception as exc:
            with self._lock:
                current = self._jobs.get(job_id)
                if current is None:
                    return
                current.status = "failed"
                current.finished_at = datetime.now(timezone.utc)
                current.error_code = "ocr_inference_failed"
                current.error_message = str(exc)
        finally:
            upload_path = Path(record.upload_path)
            upload_path.unlink(missing_ok=True)

    def _call_ocr_service(self, record: OcrJobRecord) -> tuple[str, Any]:
        endpoint = settings.OCR_SERVICE_URL.rstrip("/")
        if not endpoint:
            raise OcrJobError(
                "ocr_service_unreachable", "OCR service URL not configured"
            )

        data: dict[str, str] = {}
        if record.page_range:
            data["page_range"] = record.page_range
        if record.note:
            data["note"] = record.note
        data["output_format"] = "md_json"

        try:
            with open(record.upload_path, "rb") as file_obj:
                resp = requests.post(
                    f"{endpoint}/v1/ocr/parse",
                    files={"file": (record.original_filename, file_obj)},
                    data=data,
                    timeout=settings.OCR_JOB_TIMEOUT_SECONDS,
                )
        except requests.Timeout as exc:
            raise OcrJobError("ocr_timeout", "OCR request timed out") from exc
        except requests.RequestException as exc:
            raise OcrJobError("ocr_service_unreachable", str(exc)) from exc

        if resp.status_code != 200:
            code, message = _extract_upstream_error(resp)
            raise OcrJobError(code, message)

        try:
            payload = resp.json()
        except ValueError as exc:
            raise OcrJobError(
                "ocr_inference_failed", "Invalid OCR response payload"
            ) from exc

        markdown_text = payload.get("markdown_text")
        if markdown_text is None:
            markdown_text = payload.get("markdown", "")
        if not isinstance(markdown_text, str):
            markdown_text = str(markdown_text)

        json_data = payload.get("json_data")
        if json_data is None:
            json_data = payload.get("json")
        if json_data is None:
            json_data = {}
        return markdown_text, json_data


def _extract_upstream_error(resp: requests.Response) -> tuple[str, str]:
    message = f"OCR service returned status {resp.status_code}"
    code = "ocr_inference_failed"
    try:
        payload = resp.json()
        if isinstance(payload, dict):
            detail = (
                payload.get("detail") or payload.get("error") or payload.get("message")
            )
            if detail:
                message = str(detail)
            error_code = payload.get("error_code") or payload.get("error")
            if error_code == "oom":
                code = "oom"
            elif error_code:
                code = "ocr_inference_failed"
    except ValueError:
        if resp.text:
            message = resp.text
    return code, message


def _collect_image_regions(json_data: Any) -> list[dict[str, Any]]:
    regions: list[dict[str, Any]] = []

    def visit(node: Any, page_number: int) -> None:
        if isinstance(node, dict):
            bbox = node.get("bbox_2d")
            if _is_bbox(bbox):
                bbox_list = bbox if isinstance(bbox, list) else []
                bbox_values = [int(v) for v in bbox_list[:4]]
                label = str(node.get("label") or node.get("native_label") or "").lower()
                if _is_image_like_label(label):
                    regions.append(
                        {
                            "page_number": page_number,
                            "bbox": bbox_values,
                            "label": label or "image",
                        }
                    )
            for value in node.values():
                visit(value, page_number)
            return

        if isinstance(node, list):
            for item in node:
                visit(item, page_number)

    if isinstance(json_data, dict) and isinstance(json_data.get("pages"), list):
        for page_obj in json_data.get("pages", []):
            if not isinstance(page_obj, dict):
                continue
            page_number = int(page_obj.get("page_number") or 1)
            visit(page_obj.get("data"), page_number)
    else:
        visit(json_data, 1)

    return regions


def _is_bbox(value: Any) -> bool:
    return (
        isinstance(value, list)
        and len(value) >= 4
        and all(isinstance(item, (int, float)) for item in value[:4])
    )


def _is_image_like_label(label: str) -> bool:
    tokens = ("image", "figure", "chart", "插图", "图片", "图像")
    text = label.lower()
    return any(token in text for token in tokens)


def _sanitize_bbox(
    bbox: list[int], width: int, height: int
) -> tuple[int, int, int, int]:
    x1, y1, x2, y2 = [int(v) for v in bbox[:4]]
    if x2 < x1:
        x1, x2 = x2, x1
    if y2 < y1:
        y1, y2 = y2, y1
    x1 = max(0, min(x1, width - 1))
    y1 = max(0, min(y1, height - 1))
    x2 = max(1, min(x2, width))
    y2 = max(1, min(y2, height))
    return x1, y1, x2, y2


def _render_source_pages(
    source_path: Path, page_numbers: list[int]
) -> dict[int, Image.Image]:
    pages: dict[int, Image.Image] = {}
    ext = source_path.suffix.lower()
    if ext == ".pdf":
        document = pdfium.PdfDocument(str(source_path))
        try:
            total = len(document)
            for page_number in page_numbers:
                page_index = page_number - 1
                if page_index < 0 or page_index >= total:
                    continue
                page = document[page_index]
                bitmap = page.render(scale=2.0)
                image = bitmap.to_pil().convert("RGB")
                pages[page_number] = image
                page.close()
        finally:
            document.close()
    else:
        image = Image.open(source_path).convert("RGB")
        pages[1] = image
    return pages


ocr_job_manager = OcrJobManager()


def cleanup_expired_ocr_jobs(retention_hours: int) -> int:
    return ocr_job_manager.cleanup_expired(retention_hours)
